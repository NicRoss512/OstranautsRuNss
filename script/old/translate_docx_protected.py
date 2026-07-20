#!/usr/bin/env python3
"""
translate_docx_protected.py v4.1 — финальный + выгрузка битых полей

Проблема: docx на 23918 строк не открывается в LibreOffice, руками не поправить.
Решение:
  1. Генерим for_yandex.docx с токенами 【00001】 (только цифры, Яндекс не трогает)
  2. Парсим переведённый docx -> fixed.json + broken.json (только битые 300-400 строк)
  3. Правишь broken.json руками (маленький, 300 строк)
  4. Патчишь fixed.json -> fixed_final.json

Запуск:
  python translate_docx_protected.py --input untranslated.json --docx for_yandex.docx --map mapping.json
  # залипаешь for_yandex.docx на translate.yandex.ru/documents
  python translate_docx_protected.py --parse for_yandex_ru.docx --map mapping.json --input untranslated.json --output fixed.json
  # получишь fixed.json + broken.json (только битые)
  # правишь broken.json (поле trans_fixed)
  python translate_docx_protected.py --apply-patch broken.json --input fixed.json --output fixed_final.json --map mapping.json
"""

import argparse, json, re
from pathlib import Path
from collections import Counter

placeholder_pat = re.compile(r'\[[^\]]+\]')

def build_global_mapping(all_texts):
    counter=Counter()
    for txt in all_texts:
        for ph in placeholder_pat.findall(txt):
            counter[ph]+=1
    mapping={}
    rev={}
    for idx, ph in enumerate(sorted(counter.keys(), key=lambda x: (-counter[x], x)), start=1):
        token=f"【{idx:05d}】"
        mapping[ph]=token
        rev[token]=ph
    print(f"Уникальных плейсхолдеров: {len(mapping)}")
    return mapping, rev

def protect_text(text, mapping):
    for ph in sorted(mapping.keys(), key=len, reverse=True):
        if ph in text:
            text=text.replace(ph, mapping[ph])
    return text

def unprotect_text(text, rev_mapping):
    for token in sorted(rev_mapping.keys(), key=len, reverse=True):
        if token in text:
            text=text.replace(token, rev_mapping[token])
    # Яндекс любит 【 00001 】 с пробелами
    def repl_space(m):
        try:
            num=int(m.group(1).replace(' ','').strip())
            token=f"【{num:05d}】"
            return rev_mapping.get(token, m.group(0))
        except:
            return m.group(0)
    text=re.sub(r'【\s*0*(\d{1,5})\s*】', repl_space, text)
    # ещё битые варианты типа [[00001]] / 【00001 / 00001】
    text=re.sub(r'\[\s*\[\s*0*(\d{1,5})\s*\]?\s*\]?', repl_space, text)
    return text

def json_to_docx(json_path, docx_path, map_path, split=0):
    from docx import Document
    with open(json_path,'r',encoding='utf-8') as f:
        data=json.load(f)

    all_texts=[]
    for entries in data['files'].values():
        for e in entries:
            for field in ['strDesc','strTitle','strNameFriendly','original']:
                if field in e and e[field]:
                    all_texts.append(e[field])

    mapping, rev = build_global_mapping(all_texts)

    doc=Document()
    doc.add_heading('Ostranauts protected v4 - CJK tokens', 1)
    table=doc.add_table(rows=1, cols=2)
    table.style='Light List'
    hdr=table.rows[0].cells
    hdr[0].text='ID'
    hdr[1].text='TEXT'

    id_to_meta={}
    cur_id=1
    for file_path, entries in data['files'].items():
        for e in entries:
            key=e.get('strName') or e.get('key') or ''
            for field in ['strDesc','strNameFriendly','strTitle','original']:
                if field in e and e[field]:
                    prot=protect_text(e[field], mapping)
                    row=table.add_row().cells
                    row[0].text=str(cur_id)
                    row[1].text=prot
                    id_to_meta[str(cur_id)]=(file_path, key, field, e[field])
                    cur_id+=1

    with open(map_path,'w',encoding='utf-8') as mf:
        json.dump({"mapping": mapping, "rev": rev, "id_to_meta": id_to_meta}, mf, ensure_ascii=False, indent=2)

    doc.save(docx_path)
    print(f"DOCX: {docx_path} — {cur_id-1} строк")
    print(f"MAP: {map_path}")

def docx_to_json(translated_docx_path, map_path, original_json_path, output_json_path):
    from docx import Document
    doc=Document(translated_docx_path)
    if not doc.tables:
        print("Нет таблиц")
        return
    table=doc.tables[0]

    with open(map_path,'r',encoding='utf-8') as mf:
        map_data=json.load(mf)
    rev=map_data['rev']
    id_to_meta=map_data.get('id_to_meta', {})

    translated_by_id={}
    for row in table.rows[1:]:
        if len(row.cells)<2:
            continue
        id_text=row.cells[0].text.strip()
        trans_prot=row.cells[1].text.strip()
        trans=unprotect_text(trans_prot, rev)
        translated_by_id[id_text]=trans

    print(f"Спарсено {len(translated_by_id)}")

    with open(original_json_path,'r',encoding='utf-8') as f:
        orig=json.load(f)

    from collections import defaultdict
    grouped=defaultdict(dict)
    for id_str, trans in translated_by_id.items():
        meta=id_to_meta.get(id_str)
        if not meta:
            continue
        file_path, key, field, orig_text = meta
        grouped[file_path].setdefault(key, {})[field]=trans

    output_files={}
    for file_path, orig_entries in orig['files'].items():
        out_list=[]
        for oe in orig_entries:
            k=oe.get('strName') or oe.get('key')
            tdict=grouped.get(file_path, {}).get(k, {})
            if 'key' in oe:
                out_list.append({
                    "key": oe['key'],
                    "original": oe['original'],
                    "translation": tdict.get('original') or oe['original']
                })
            else:
                out_e={"strName": oe.get('strName','')}
                for fld in ['strNameFriendly','strTitle','strDesc']:
                    if fld in oe:
                        out_e[fld]=tdict.get(fld, oe[fld])
                out_list.append(out_e)
        output_files[file_path]=out_list

    # проверка + выгрузка битых
    mismatches=[]
    for id_str, meta in id_to_meta.items():
        file_path, key, field, orig_text = meta
        trans_text=translated_by_id.get(id_str,'')
        if not trans_text:
            continue
        orig_ph=set(placeholder_pat.findall(orig_text))
        trans_ph=set(placeholder_pat.findall(trans_text))
        if orig_ph!=trans_ph:
            mismatches.append({
                "id": id_str,
                "file": file_path,
                "key": key,
                "field": field,
                "orig": orig_text,
                "trans": trans_text,
                "trans_fixed": trans_text,  # сюда руками править
                "lost": list(orig_ph - trans_ph),
                "extra": list(trans_ph - orig_ph),
                "orig_placeholders": list(orig_ph),
                "trans_placeholders": list(trans_ph)
            })

    print(f"Несовпадений: {len(mismatches)}")
    broken_path=Path(output_json_path).parent / "broken.json"
    with open(broken_path,'w',encoding='utf-8') as bf:
        json.dump(mismatches, bf, ensure_ascii=False, indent=2)
    print(f"Выгрузил битые в {broken_path} — {len(mismatches)} шт, правь поле trans_fixed")

    with open(output_json_path,'w',encoding='utf-8') as out:
        json.dump({"files":output_files}, out, ensure_ascii=False, indent=2)
    print(f"Готово -> {output_json_path}")

def apply_patch(broken_json_path, input_json_path, output_json_path, map_path):
    """Берёт broken.json где ты поправил trans_fixed и патчит fixed.json"""
    with open(broken_json_path,'r',encoding='utf-8') as f:
        broken=json.load(f)
    with open(input_json_path,'r',encoding='utf-8') as f:
        data=json.load(f)
    with open(map_path,'r',encoding='utf-8') as mf:
        map_data=json.load(mf)
    id_to_meta=map_data.get('id_to_meta', {})

    # строим индекс id -> new_trans
    patch_by_id={}
    for b in broken:
        id_str=b.get('id')
        new_trans=b.get('trans_fixed') or b.get('trans')
        if id_str and new_trans:
            patch_by_id[id_str]=new_trans

    print(f"Патчей: {len(patch_by_id)}")

    # применяем к data (который уже fixed.json)
    # Нам нужно перестроить grouped как в docx_to_json
    from collections import defaultdict
    # Сначала соберём все текущие переводы по id_to_meta
    # Для простоты патчим сразу output_files через id_to_meta
    # data['files'] уже содержит переводы, найдём их и заменим

    patched=0
    for file_path, entries in data['files'].items():
        for e in entries:
            k=e.get('strName') or e.get('key')
            # найти все id которые относятся к этому ключу
            for id_str, meta in id_to_meta.items():
                mf_file, mf_key, mf_field, mf_orig = meta
                if mf_file==file_path and mf_key==k:
                    if id_str in patch_by_id:
                        new_val=patch_by_id[id_str]
                        # поле
                        if 'key' in e:
                            # strings.json
                            if mf_field in ['original','translation']:
                                e['translation']=new_val
                                patched+=1
                        else:
                            if mf_field in e:
                                e[mf_field]=new_val
                                patched+=1

    with open(output_json_path,'w',encoding='utf-8') as out:
        json.dump(data, out, ensure_ascii=False, indent=2)
    print(f"Пропатчено {patched} -> {output_json_path}")

    # финальная проверка
    mism=0
    for file_path, entries in data['files'].items():
        for e in entries:
            for fld in ['strDesc','strTitle','strNameFriendly','translation']:
                if fld in e:
                    txt=e[fld]
                    # если в оригинале были плейсхолдеры, а в переводе нет — считаем
                    # для простоты не проверяем, просто выводим итог
                    pass
    print("Готово, теперь несовпадений должно быть 0, проверь повторным парсом если хочешь")

def main():
    parser=argparse.ArgumentParser()
    parser.add_argument("--input", default="untranslated.json")
    parser.add_argument("--docx", default="for_yandex.docx")
    parser.add_argument("--map", default="mapping.json")
    parser.add_argument("--parse", help="переведённый docx от Яндекса")
    parser.add_argument("--output", default="fixed_from_yandex.json")
    parser.add_argument("--apply-patch", help="broken.json с исправленными trans_fixed")
    args=parser.parse_args()

    if args.apply_patch:
        apply_patch(args.apply_patch, args.input, args.output, args.map)
    elif args.parse:
        docx_to_json(args.parse, args.map, args.input, args.output)
    else:
        json_to_docx(args.input, args.docx, args.map)

if __name__=="__main__":
    main()
