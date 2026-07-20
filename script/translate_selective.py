#!/usr/bin/env python3
"""
Экспорт:
  python translate_selective.py --export --str strTitle,strDesc -i data/interactions/interactions.json -o interactions.docx --dict dict.json
  # создаст interactions.docx с 1 колонкой TEXT и mapping.json с привязкой по порядку

Импорт:
  python translate_selective.py --import -i interactions_translated.docx --map mapping.json -o interactions_fixed.json --dict dict.json --original data/interactions/interactions.json

dict.json — внешний словарь [us]->Иван чтобы Яндекс видел контекст, на импорте Иван->[us]
"""

import argparse, json, re
from pathlib import Path

placeholder_pat = re.compile(r'\[[^\]]+\]')

DEFAULT_DICT = {
    "[us]": "ИванЮС",
    "[them]": "ПётрЗЭМ",
    "[3rd]": "СидорТРТ",
    "[us-pos]": "ИванПОС",
    "[them-pos]": "ПётрПОС",
    "[us-subj]": "ИванСАБ",
    "[them-subj]": "ПётрСАБ",
    "[us-obj]": "ИванОБЖ",
    "[them-obj]": "ПётрОБЖ",
    "[3rd-pos]": "СидорПОС",
    "[3rd-subj]": "СидорСАБ",
    "[3rd-obj]": "СидорОБЖ",
    "[us-reflexive]": "ИванРЕФ",
    "[is]": "",
    "[has]": "",
    "[was]": "былБЫЛ",
    "[were]": "былиБЫЛИ",
    "[gains]": "получаетГЕЙНС",
    "[says]": "говоритСЕЙС",
}

def load_dict(dict_path):
    if dict_path and Path(dict_path).exists():
        try:
            with open(dict_path,'r',encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    # создаём пример
    if dict_path:
        with open(dict_path,'w',encoding='utf-8') as out:
            json.dump(DEFAULT_DICT, out, ensure_ascii=False, indent=2)
        print(f"Создал пример словаря {dict_path}")
    return DEFAULT_DICT.copy()

def protect_with_dict(text, mapping):
    # Заменяем по длине убывания
    for ph in sorted(mapping.keys(), key=len, reverse=True):
        if ph in text:
            # Если маппинг пустой — заменяем на пробел-заглушку чтобы потом склеить
            repl=mapping[ph]
            text=text.replace(ph, repl)
    # Если были пустые замены типа [is]->"" то остались двойные пробелы: "Иван  стоит" -> "Иван стоит"
    # Склеиваем 2+ пробелов в один, но сохраняем \n
    text=re.sub(r'[ \t]{2,}', ' ', text)
    text=re.sub(r' \n', '\n', text)
    text=re.sub(r'\n ', '\n', text)
    # Убираем пробел перед знаками препинания если вдруг: "Иван ,"
    text=re.sub(r'\s+([,.!?:;])', r'\1', text)
    return text.strip()

def unprotect_with_dict(text, rev_mapping):
    # rev_mapping: "Иван" -> "[us]"
    # Сортируем по длине русского слова чтобы "его" не заменило часть "его-то"
    for word, ph in sorted(rev_mapping.items(), key=lambda x: len(x[0]), reverse=True):
        if not word:
            continue
        # Заменяем только целые слова? Для имён типа Иван — только если окружено не буквами
        # Используем простую замену, но с границами для коротких слов типа "он", "его"
        if len(word)<=3:
            # для коротких местоимений требуем границы
            pattern=r'(?<![А-Яа-яЁёA-Za-z])' + re.escape(word) + r'(?![А-Яа-яЁёA-Za-z])'
            text=re.sub(pattern, ph, text)
        else:
            if word in text:
                text=text.replace(word, ph)
    return text

def collect_placeholders_from_texts(files_dict, fields):
    """Собирает все [xxx] из выбранных полей"""
    found=set()
    for entries in files_dict.values():
        if isinstance(entries, dict):
            entries=[entries]
        for e in entries:
            if not isinstance(e, dict):
                continue
            for field in fields:
                if field in e and e[field]:
                    for ph in placeholder_pat.findall(e[field]):
                        found.add(ph)
    return found

def has_russian(text):
    """Есть ли кириллица"""
    return bool(re.search(r'[А-Яа-яЁё]', text))

def load_input_json(input_path, file_filter=None):
    with open(input_path,'r',encoding='utf-8') as f:
        data=json.load(f)
    if isinstance(data, dict) and 'files' in data:
        if file_filter:
            for fp in data['files']:
                if file_filter in fp:
                    print(f"Фильтр {file_filter} -> {fp}")
                    return {fp: data['files'][fp]}, data
        return data['files'], data
    elif isinstance(data, list):
        return {str(input_path): data}, {"files": {str(input_path): data}}
    else:
        return {str(input_path): [data]}, {"files": {str(input_path): [data]}}

def export_docx(args):
    dict_map=load_dict(args.dict)
    files_dict, _ = load_input_json(args.input, args.file_filter)
    fields=[s.strip() for s in args.str_fields.split(',') if s.strip()]
    print(f"Поля: {fields}, словарь до: {len(dict_map)}")

    # Авто-добавление плейсхолдеров которых нет в dict.json
    found_ph=collect_placeholders_from_texts(files_dict, fields)
    added_list=[]
    for ph in found_ph:
        if ph not in dict_map:
            # дефолт: для вспомогательных [is],[has] -> пусто чтобы не было двойных пробелов, для остальных -> без скобок
            if ph in ["[is]","[has]","[is-not]","[does]","[does-not]","[did]"]:
                dict_map[ph]=""
            elif ph in DEFAULT_DICT:
                dict_map[ph]=DEFAULT_DICT[ph]
            else:
                # глагольные типа [gains] -> gains без скобок
                inner=ph.strip('[]')
                dict_map[ph]=inner
            added_list.append(ph)
    if added_list:
        print(f"Авто-добавил {len(added_list)} новых плейсхолдеров в словарь: {sorted(added_list)[:20]}")
        # сохраняем обновлённый словарь обратно
        with open(args.dict,'w',encoding='utf-8') as out:
            json.dump(dict_map, out, ensure_ascii=False, indent=2)
        print(f"Обновил {args.dict} — теперь {len(dict_map)} записей, проверь что там")

    print(f"Словарь после: {len(dict_map)} (пустых [is]->'': {sum(1 for v in dict_map.values() if not v)}) — плейсхолдеры: {sorted(found_ph)}")

    # Собираем все строки для экспорта
    id_to_meta={}
    all_protected=[]
    cur_id=0
    skipped_russian=0
    for file_path, entries in files_dict.items():
        if isinstance(entries, dict):
            entries=[entries]
        for e in entries:
            if not isinstance(e, dict):
                continue
            key=e.get('strName') or e.get('key') or ''
            for field in fields:
                if field in e and e[field]:
                    orig=e[field]
                    # Фильтр: только непереведённые (без кириллицы)
                    if getattr(args, 'only_untranslated', False) and has_russian(orig):
                        skipped_russian+=1
                        continue
                    prot=protect_with_dict(orig, dict_map)
                    all_protected.append(prot)
                    id_to_meta[str(cur_id)]=(file_path, key, field, orig)
                    cur_id+=1

    if getattr(args, 'only_untranslated', False):
        print(f"Пропущено уже переведённых (с кириллицей): {skipped_russian}")

    # Сохраняем маппинг
    with open(args.map,'w',encoding='utf-8') as mf:
        json.dump({"dict": dict_map, "rev": {v:k for k,v in dict_map.items() if v}, "id_to_meta": id_to_meta, "fields": fields}, mf, ensure_ascii=False, indent=2)

    # Определяем формат по расширению выходного файла
    out_path=Path(args.output)
    if out_path.suffix.lower()=='.txt':
        # Текстовый формат: каждая строка — одна строка для перевода
        # Пустые строки сохраняем как есть? Пишем с сохранением \n внутри строки заменяя на \\n ?
        # Для простоты — каждая запись в отдельной строке, \n внутри экранируем как ⏎ или оставляем literal \n ?
        # Будем писать каждую запись на отдельной строке, заменяя реальные \n на \n литерал \n -> \\n
        # Чтобы при импорте восстановить
        with open(args.output,'w',encoding='utf-8') as out_f:
            for line in all_protected:
                # Экранируем \n чтобы не ломать построчность
                out_f.write(line.replace('\r','').replace('\n','\\n') + '\n')
        print(f"Экспорт TXT: {args.output} — {cur_id} строк, map {args.map}")
    else:
        # DOCX формат как раньше
        from docx import Document
        doc=Document()
        table=doc.add_table(rows=0, cols=1)
        table.style='Table Grid'
        for prot in all_protected:
            row=table.add_row().cells
            row[0].text=prot
        doc.save(args.output)
        print(f"Экспорт DOCX: {args.output} — {cur_id} строк (только текст), map {args.map}")

def import_docx(args):
    with open(args.map,'r',encoding='utf-8') as mf:
        map_data=json.load(mf)
    dict_map=map_data.get('dict', {})
    if args.dict and Path(args.dict).exists():
        dict_map=load_dict(args.dict)
    rev_map={v:k for k,v in dict_map.items() if v}
    id_to_meta=map_data.get('id_to_meta', {})

    # Читаем входной файл — может быть docx или txt
    translated_texts=[]
    input_path=Path(args.input)
    if input_path.suffix.lower()=='.txt':
        # Текстовый формат: каждая строка — одна запись, \n внутри экранирован как \\n
        with open(args.input,'r',encoding='utf-8') as f:
            for line in f:
                # Убираем только \n в конце строки, но сохраняем пустые?
                txt=line.rstrip('\n\r')
                # Декодируем \\n обратно в \n
                txt=txt.replace('\\n','\n')
                if txt or True:  # даже пустые считаем? Пропускаем пустые чтобы не сдвинуть нумерацию? Лучше пропускать пустые как в docx версии
                    # В экспорте мы пропускали пустые? Нет, мы писали все
                    # Для txt будем брать все строки, даже пустые? Пропустим пустые как раньше
                    if txt.strip()=='':
                        # если пустая строка — это может быть пустой перевод, считаем как пустую
                        # но чтобы не сдвинуть нумерацию, лучше не пропускать? В docx версии мы пропускали пустые
                        # Для консистентности — пропускаем пустые как в docx версии
                        continue
                    translated_texts.append(txt)
    else:
        from docx import Document
        doc=Document(args.input)
        # читаем все параграфы/ячейки по порядку — только текст, без ID внутри
        if doc.tables:
            # если таблица 1 колонка
            table=doc.tables[0]
            for row in table.rows:
                txt=row.cells[0].text.strip()
                if not txt:
                    continue
                translated_texts.append(txt)
        else:
            for para in doc.paragraphs:
                txt=para.text.strip()
                if txt:
                    translated_texts.append(txt)

        # если в таблице была строка заголовка — она попадёт, но id_to_meta начинается с 0
        if len(translated_texts)==len(id_to_meta)+1:
            if "TEXT" in translated_texts[0] or "Export" in translated_texts[0]:
                translated_texts=translated_texts[1:]

    print(f"Спарсено {len(translated_texts)} строк из {args.input}, ожидалось {len(id_to_meta)}")

    # восстанавливаем плейсхолдеры
    restored=[]
    for txt in translated_texts:
        restored.append(unprotect_with_dict(txt, rev_map))

    # теперь маппим по порядку (боксам)
    from collections import defaultdict
    grouped=defaultdict(dict)
    for idx, trans in enumerate(restored):
        id_str=str(idx)
        meta=id_to_meta.get(id_str)
        if not meta:
            # если docx был с заголовком — пробуем id_str+1
            meta=id_to_meta.get(str(idx+1))
            if not meta:
                continue
        file_path, key, field, orig = meta
        grouped[file_path].setdefault(key, {})[field]=trans

    # загружаем оригинал для структуры
    orig_path=args.original
    if not orig_path:
        # берём первый file_path из id_to_meta
        if id_to_meta:
            first_fp=list(id_to_meta.values())[0][0]
            if Path(first_fp).exists():
                orig_path=first_fp

    if orig_path and Path(orig_path).exists():
        with open(orig_path,'r',encoding='utf-8') as f:
            orig_json=json.load(f)

        if isinstance(orig_json, dict) and 'files' in orig_json:
            for fp, entries in orig_json['files'].items():
                for e in entries:
                    k=e.get('strName') or e.get('key')
                    if fp in grouped and k in grouped[fp]:
                        for fld, new_val in grouped[fp][k].items():
                            if fld in e:
                                e[fld]=new_val
            with open(args.output,'w',encoding='utf-8') as out:
                json.dump(orig_json, out, ensure_ascii=False, indent=2)
            print(f"Пропатчен большой JSON -> {args.output}")
        elif isinstance(orig_json, list):
            for e in orig_json:
                k=e.get('strName') or e.get('key')
                # ищем в grouped по любому файлу
                for fp in grouped:
                    if k in grouped[fp]:
                        for fld, new_val in grouped[fp][k].items():
                            if fld in e:
                                e[fld]=new_val
            with open(args.output,'w',encoding='utf-8') as out:
                json.dump(orig_json, out, ensure_ascii=False, indent=2)
            print(f"Готово single file -> {args.output}")
        else:
            with open(args.output,'w',encoding='utf-8') as out:
                json.dump(grouped, out, ensure_ascii=False, indent=2)
    else:
        with open(args.output,'w',encoding='utf-8') as out:
            json.dump(grouped, out, ensure_ascii=False, indent=2)
        print(f"Сохранён grouped -> {args.output}")

    # проверка
    mism=0
    for id_str, meta in id_to_meta.items():
        file_path, key, field, orig_text = meta
        idx=int(id_str)
        if idx>=len(restored):
            continue
        trans=restored[idx]
        if set(placeholder_pat.findall(orig_text))!=set(placeholder_pat.findall(trans)):
            mism+=1
            if mism<=10:
                print(f"MISMATCH ID={id_str} {file_path} {key} {field} ORIG {set(placeholder_pat.findall(orig_text))} -> TRANS {set(placeholder_pat.findall(trans))}")
    print(f"Несовпадений: {mism}")

def main():
    parser=argparse.ArgumentParser(description="Только строки, без путей внутри docx")
    sub=parser.add_mutually_exclusive_group(required=True)
    sub.add_argument('--export', action='store_true')
    sub.add_argument('--import', dest='import_mode', action='store_true')
    parser.add_argument('-i','--input', required=True)
    parser.add_argument('-o','--output', required=True)
    parser.add_argument('--str', dest='str_fields', default='strTitle,strDesc')
    parser.add_argument('-f','--file-filter', dest='file_filter')
    parser.add_argument('--dict', default='dict.json')
    parser.add_argument('--map', default='mapping.json')
    parser.add_argument('--original', help='оригинальный json для патча')
    parser.add_argument('--only-untranslated', action='store_true', help='вытаскивать только строки без русских букв (непереведённые)')
    args=parser.parse_args()

    if args.export:
        export_docx(args)
    else:
        import_docx(args)

if __name__=="__main__":
    main()
